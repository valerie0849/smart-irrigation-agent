import fitz
import pdfplumber
import os
from typing import List, Dict, Any

class DocumentParser:
    def __init__(self):
        pass
    
    def parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        try:
            return self._parse_with_pdfplumber(file_path)
        except:
            try:
                return self._parse_with_pymupdf(file_path)
            except:
                return []
    
    def _parse_with_pdfplumber(self, file_path: str) -> List[Dict[str, Any]]:
        all_text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    all_text += text + "\n\n"
        if all_text:
            return [{
                "page": 1,
                "content": all_text.strip(),
                "source": os.path.basename(file_path),
                "chunk_id": "full_doc"
            }]
        return []
    
    def _parse_with_pymupdf(self, file_path: str) -> List[Dict[str, Any]]:
        all_text = ""
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            if text:
                all_text += text + "\n\n"
        if all_text:
            return [{
                "page": 1,
                "content": all_text.strip(),
                "source": os.path.basename(file_path),
                "chunk_id": "full_doc"
            }]
        return []
    
    def parse_text(self, file_path: str) -> List[Dict[str, Any]]:
        chunks = []
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            chunks.append({
                "page": 1,
                "content": content,
                "source": os.path.basename(file_path),
                "chunk_id": "full_text"
            })
        return chunks

    def parse_word(self, file_path: str) -> List[Dict[str, Any]]:
        import docx
        doc = docx.Document(file_path)
        all_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                all_text.append(" | ".join(cells))
        content = "\n".join(all_text)
        if content:
            return [{
                "page": 1,
                "content": content,
                "source": os.path.basename(file_path),
                "chunk_id": "full_doc"
            }]
        return []

    def parse_document(self, file_path: str) -> List[Dict[str, Any]]:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return self.parse_pdf(file_path)
        elif ext in ('.txt', '.md'):
            return self.parse_text(file_path)
        elif ext in ('.doc', '.docx'):
            return self.parse_word(file_path)
        else:
            return []