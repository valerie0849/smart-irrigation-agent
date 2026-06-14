import os
os.environ["OMP_NUM_THREADS"] = "1"

from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse, Response
from sqlalchemy.orm import Session
from backend.src.models import get_db, init_db, ConversationHistory
from backend.agents.requirement import RequirementAgent
from backend.agents.irrigation_design import IrrigationDesignAgent
from backend.agents.equipment import EquipmentAgent
from backend.agents.orchestrator import OrchestratorAgent
from backend.agents.quality import QualityAgent
from backend.knowledge_forest.main import KnowledgeForest
from backend.data_acquisition.main import DataAcquisition, CHINA_COORDS
from backend.agents.requirement import _infer_crop_by_location
from backend.llm import deepseek, build_irrigation_prompt, build_quality_prompt, build_conversation_reply, log_json
from pydantic import BaseModel
from typing import Dict, Any, Optional, List, AsyncGenerator
import logging
import uuid
import os
import json
import io
import asyncio

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

frontend_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "frontend")

app = FastAPI(title="智慧灌溉方案智能生成系统", version="3.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

knowledge_forest = None
conversation_context: Dict[str, List[Dict]] = {}
session_entities: Dict[str, Dict[str, Any]] = {}

class ChatMsg(BaseModel):
    session_id: Optional[str] = None
    message: str
    original_plan: Optional[str] = None

@app.on_event("startup")
async def startup():
    global knowledge_forest, _pending_incremental
    logger.info("=" * 60)
    logger.info("[启动] 智慧灌溉方案智能生成系统")
    init_db()
    db = next(get_db())
    try:
        knowledge_forest = KnowledgeForest(db)
        if knowledge_forest.try_load():
            logger.info("[启动] 知识森林从缓存加载")
        else:
            logger.info("[启动] 知识森林需首次构建")
    finally:
        db.close()

    if knowledge_forest and knowledge_forest.is_built:
        _pending_incremental = True
    deepseek.load()
    logger.info("[启动] 系统就绪")

@app.get("/health")
async def health():
    return {"status": "healthy", "llm": deepseek._loaded, "kf": knowledge_forest is not None and knowledge_forest.is_built}

CROP_NAMES = ["小麦", "水稻", "玉米", "棉花", "蔬菜", "果树", "大豆", "花生", "油菜"]
GEN_TRIGGER_WORDS = ["生成方案", "设计方案", "制定方案", "出方案", "帮我设计", "帮我生成", "帮我做方案", "来一份方案", "给我一份方案"]
MODIFY_TRIGGER_WORDS = ["修改方案", "对已有方案提出修改", "调整方案", "修改已有方案", "优化方案", "改进方案", "帮我把", "参数调整", "变更方案", "调整一下", "改一下", "换一下", "修改一下", "方案修改", "方案调整"]
KNOWLEDGE_TRIGGER_WORDS = ["知识", "依据", "标准", "规范", "规定", "参考标准", "技术标准", "参考文献"]
KNOWLEDGE_BROWSE_WORDS = ["查看知识森林", "知识森林", "知识森林概况", "知识库概况", "知识列表", "所有知识", "知识全貌", "浏览知识", "列出知识", "知识库有哪些", "查看知识库", "查看知识", "查询知识库", "查询知识"]
WEATHER_TRIGGER_WORDS = ["天气", "气象", "气温", "降水", "降雨", "湿度"]
SOIL_TRIGGER_WORDS = ["土壤"]


def _classify_intent(text: str) -> str:
    u = text.lower()

    if any(kw in u for kw in MODIFY_TRIGGER_WORDS):
        return "modify"

    if any(kw in u for kw in WEATHER_TRIGGER_WORDS):
        return "weather"

    if any(kw in u for kw in SOIL_TRIGGER_WORDS):
        return "soil"

    if any(kw in u for kw in KNOWLEDGE_BROWSE_WORDS):
        return "knowledge_browse"

    if any(kw in u for kw in KNOWLEDGE_TRIGGER_WORDS):
        return "knowledge"

    if any(kw in u for kw in GEN_TRIGGER_WORDS):
        return "generate"

    has_area = any(kw in text for kw in ["亩", "公顷"])
    has_crop = any(kw in text for kw in CROP_NAMES)
    has_budget = any(kw in text for kw in ["预算", "万"])
    has_location = any(loc in text for loc in CHINA_COORDS.keys() if len(loc) > 0)

    if has_location and has_crop:
        return "generate"

    if has_location and (has_area or has_budget):
        return "generate"

    if (has_area and has_crop) or (has_area and has_budget) or (has_crop and has_location and has_area):
        return "generate"

    if has_crop and has_area:
        return "generate"

    if has_area or has_budget:
        return "generate"

    return "unknown"


def _analyze_missing_info(text: str, accumulated: dict = None) -> List[str]:
    entities = RequirementAgent._raw_extract_entities(text)
    if accumulated:
        for k, v in accumulated.items():
            if v and not entities.get(k):
                entities[k] = v
    missing = []
    if not entities.get("location"):
        missing.append("项目地点（如：河南中牟县、山东菏泽）")
    if not entities.get("crop_type") and not entities.get("location"):
        missing.append("作物类型（如：小麦、水稻、玉米、棉花）")
    if not entities.get("area"):
        missing.append("灌溉面积（如：2000亩）")
    if not entities.get("budget"):
        missing.append("预算范围（如：预算200万）")
    return missing, entities

def _save_msg_to_db(db, sid, role, content):
    try:
        h = ConversationHistory(session_id=sid, role=role, content=content, context={})
        db.add(h)
        db.commit()
    except Exception as e:
        logger.error(f"[DB] 保存消息失败: {e}")
        db.rollback()


def _load_history_from_db(db, sid):
    try:
        rows = db.query(ConversationHistory).filter(ConversationHistory.session_id == sid).order_by(ConversationHistory.timestamp).all()
        return [{"role": r.role, "content": r.content} for r in rows]
    except Exception as e:
        logger.error(f"[DB] 加载历史失败: {e}")
        return []


@app.post("/chat")
async def chat(msg: ChatMsg, db: Session = Depends(get_db)):
    global knowledge_forest, conversation_context
    sid = msg.session_id or str(uuid.uuid4())
    ui = msg.message.strip()

    logger.info(f"[对话] sid={sid[:8]} msg={ui[:80]}")
    log_json("用户消息", {"session_id": sid, "message": ui})

    if sid not in conversation_context:
        conversation_context[sid] = _load_history_from_db(db, sid)
    conversation_context[sid].append({"role": "user", "content": ui})
    _save_msg_to_db(db, sid, "user", ui)
    ctx_history = conversation_context[sid]

    intent = _classify_intent(ui)
    logger.info(f"[路由] intent={intent} msg={ui[:50]}")

    if sid not in session_entities:
        session_entities[sid] = {}
    new_ents = RequirementAgent._raw_extract_entities(ui)
    for k, v in new_ents.items():
        if v:
            session_entities[sid][k] = v
    accumulated = session_entities[sid]
    logger.info(f"[实体] 本轮: {new_ents}, 累计: {accumulated}")

    if intent == "generate":
        has_plan = any(
            m.get("role") == "assistant"
            and len(m.get("content", "")) > 500
            and any(ch in m.get("content", "") for ch in ["一、", "1. 项目概况", "二、", "三、"])
            for m in ctx_history
        )
        is_param_tweak = (
            not new_ents.get("location") and
            not new_ents.get("crop_type") and
            not new_ents.get("area")
        ) and (
            new_ents.get("budget") or
            any(kw in ui for kw in ["预算", "万", "修改", "调整", "改成", "改为", "换成", "变成", "更换", "替换", "更新", "变更", "换成", "改一下", "调一下"])
        )
        logger.info(f"[路由] has_plan={has_plan}, is_param_tweak={is_param_tweak}")
        if has_plan and is_param_tweak:
            logger.info(f"[路由] 检测到参数微调，转为modify")
            intent = "modify"

    if intent == "knowledge_browse":
        return await _do_knowledge_browse(db, sid)

    if intent == "knowledge":
        return await _do_knowledge(db, sid, ui)

    if intent == "weather":
        return await _do_weather(sid, ui)

    if intent == "soil":
        return await _do_soil(sid, ui)

    if intent == "generate":
        missing, entities = _analyze_missing_info(ui, session_entities.get(sid, {}))
        if missing:
            reply = "好的，我来帮您生成灌溉方案！不过还需要以下信息：\n\n" + "\n".join(f"- {m}" for m in missing) + f"\n\n已有信息：{json.dumps({k:v for k,v in entities.items() if v}, ensure_ascii=False)}" if len(missing) < 4 else "您提到了一些信息，但还需要补充以下关键内容：\n\n" + "\n".join(f"- {m}" for m in missing) + "\n\n请补充这些信息后，我会为您生成专业的灌溉方案。"
            ctx_history.append({"role": "assistant", "content": reply})
            _save_msg_to_db(db, sid, "assistant", reply)
            return {"success": True, "session_id": sid, "reply": reply, "type": "clarification", "missing": [m.split(" ")[1] if " " in m else m for m in missing]}

        msginfo = json.dumps({k: str(v)[:80] for k, v in entities.items()}, ensure_ascii=False)
        ctx_history.append({"role": "assistant", "content": "正在启动方案生成流程..."})
        _save_msg_to_db(db, sid, "assistant", "正在启动方案生成流程...")
        return {
            "success": True, "session_id": sid,
            "type": "generate_start",
            "message": ui,
            "entities": entities,
        }

    if intent == "modify":
        # 查找历史中的方案内容
        plan_content = None
        for msg in reversed(conversation_context[sid]):
            if msg.get("role") == "assistant" and len(msg.get("content", "")) > 500 and any(ch in msg.get("content", "") for ch in ["一、", "1. 项目概况", "二、", "三、"]):
                plan_content = msg["content"]
                break
        
        if not plan_content:
            reply = "看起来您想修改方案，但我还没有找到之前生成的方案。请先生成一份方案后再尝试修改。"
            ctx_history.append({"role": "assistant", "content": reply})
            _save_msg_to_db(db, sid, "assistant", reply)
            return {"success": True, "session_id": sid, "reply": reply, "type": "clarification"}
        
        ctx_history.append({"role": "assistant", "content": "正在启动方案修改流程..."})
        _save_msg_to_db(db, sid, "assistant", "正在启动方案修改流程...")
        return {
            "success": True, "session_id": sid,
            "type": "modify_start",
            "message": ui,
            "original_plan": plan_content,
        }

    if intent == "unknown":
        has_crop = any(kw in ui for kw in CROP_NAMES)
        has_water = any(kw in ui for kw in ["灌溉", "灌水", "滴灌", "喷灌", "微灌", "节水"])
        if has_crop or has_water:
            reply = "您似乎想了解灌溉相关信息，但还需要更多细节才能生成方案。\n\n请告诉我：\n- 项目地点\n- 🌾 作物类型\n- 📐 灌溉面积\n- 💰 预算范围\n\n例如：\"河南中牟县2000亩小麦，预算200万\""
            ctx_history.append({"role": "assistant", "content": reply})
            _save_msg_to_db(db, sid, "assistant", reply)
            return {"success": True, "session_id": sid, "reply": reply, "type": "clarification"}

    history_str = "\n".join([f"{m['role']}: {m['content'][:200]}..." if len(m['content']) > 200 else f"{m['role']}: {m['content']}" for m in ctx_history[-10:]])
    if deepseek._loaded:
        prompt = build_conversation_reply(ui, history_str, session_entities.get(sid, {}))
        llm_reply = await deepseek.generate_async(prompt, max_tokens=2048)
        reply = llm_reply if llm_reply else '你好！我是智慧灌溉助手 。\n\n你可以：\n1. 描述项目，如"河南中牟县2000亩小麦，预算200万"，我来生成方案\n2. 不指定作物时，系统会根据地区自动推断最适合的作物\n3. 查询气象/土壤数据\n4. 查询知识库\n5. 对已有方案提出修改'
    else:
        reply = '你好！我是智慧灌溉助手 。\n\n你可以：\n1. 描述项目，如"河南中牟县2000亩小麦，预算200万"，我来生成方案\n2. 不指定作物时，系统会根据地区自动推断最适合的作物\n3. 查询气象/土壤数据\n4. 查询知识库\n5. 对已有方案提出修改'
    ctx_history.append({"role": "assistant", "content": reply})
    _save_msg_to_db(db, sid, "assistant", reply)
    return {"success": True, "session_id": sid, "reply": reply, "type": "conversation"}


@app.post("/chat/generate")
async def chat_generate_stream(msg: ChatMsg, db: Session = Depends(get_db)):
    global knowledge_forest, conversation_context, session_entities
    sid = msg.session_id or str(uuid.uuid4())
    ui = msg.message.strip()

    if sid not in conversation_context:
        conversation_context[sid] = _load_history_from_db(db, sid)
    ctx_history = conversation_context[sid]

    if sid not in session_entities:
        session_entities[sid] = {}
    new_ents = RequirementAgent._raw_extract_entities(ui)
    for k, v in new_ents.items():
        if v:
            session_entities[sid][k] = v
    accumulated = session_entities[sid]

    async def event_stream() -> AsyncGenerator[str, None]:
        global knowledge_forest
        yield _sse({"type": "progress", "step": "start", "message": "开始生成灌溉方案..."})

        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "requirement", "message": "正在分析项目需求，提取关键参数..."})

        req = RequirementAgent(db)
        profile = await req.analyze(ui, accumulated)
        log_json("1-需求分析结果", profile)

        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "knowledge", "message": "正在初始化知识森林..."})

        if knowledge_forest is None:
            knowledge_forest = KnowledgeForest(db)

        try:
            if not knowledge_forest.is_built:
                if not knowledge_forest.try_load():
                    await knowledge_forest.build_forest()
        except Exception as e:
            logger.error(f"[方案生成] 知识森林初始化失败: {e}")

        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "design", "message": "正在设计灌溉制度（结合知识森林案例），计算作物需水量和灌水周期..."})

        design = IrrigationDesignAgent(db, knowledge_forest)
        irrig = await design.design(profile)
        log_json("2-灌溉制度设计", {
            "irrigation_method": irrig.get("project_info", {}).get("irrigation_method"),
            "total_water": irrig.get("project_info", {}).get("total_water_requirement_m3"),
            "knowledge_informed": irrig.get("project_info", {}).get("knowledge_informed"),
            "schedule": irrig.get("irrigation_schedule", []),
            "calc": irrig.get("calculation_details", {}),
        })

        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "equipment", "message": "正在匹配灌溉设备（参考知识森林配置），优化选型方案和成本核算..."})

        eq = EquipmentAgent(db, knowledge_forest)
        budget = profile.get("budget", 1000000)
        eqr = await eq.select(irrig, budget)
        log_json("3-设备选型结果", {
            "total_cost": eqr.get("total_cost"),
            "equipment_count": len(eqr.get("equipment_list", [])),
        })

        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "knowledge_query", "message": "正在多通道检索知识森林（模板/技术/合规三通道）..."})

        try:
            qt = f"{profile.get('crop_type','')} {profile.get('location','')} 灌溉制度设计"
            channel_results = await knowledge_forest.multi_channel_query(qt, profile, k_per_channel=5)

            template_refs = channel_results.get("template", [])
            tech_refs = channel_results.get("domain_knowledge", [])
            standard_refs = channel_results.get("standard", [])

            krefs = []
            krefs.extend([{**r, "channel": "template"} for r in template_refs[:3]])
            krefs.extend([{**r, "channel": "tech"} for r in tech_refs[:3]])
            krefs.extend([{**r, "channel": "standard"} for r in standard_refs[:3]])

            if len(krefs) < 5:
                fallback = await knowledge_forest.query(qt, k=5)
                existing_sources = {r.get("source", "") for r in krefs}
                for r in fallback:
                    if r.get("source", "") not in existing_sources and len(krefs) < 5:
                        r["channel"] = "mixed"
                        krefs.append(r)
                        existing_sources.add(r.get("source", ""))

            log_json("4-多通道知识检索", {
                "query": qt,
                "channels": {
                    "template": len(template_refs),
                    "tech": len(tech_refs),
                    "standard": len(standard_refs),
                },
                "total": len(krefs),
                "top_sources": [{"source": r["source"], "similarity": r.get("similarity", 0), "channel": r.get("channel", "")} for r in krefs[:5]],
            })
            logger.info(f"[方案生成] krefs 最终数量: {len(krefs)}, 样本: {[{'source': r.get('source'), 'channel': r.get('channel')} for r in krefs[:3]]}")
        except Exception as e:
            logger.error(f"[方案生成] 多通道检索失败: {e}")
            krefs = []
            template_refs = []
            tech_refs = []
            standard_refs = []

        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "orchestrate", "message": "正在整合方案内容，生成结构化报告..."})

        orch = OrchestratorAgent(db)
        orch_result = await orch.orchestrate({
            "project_profile": profile,
            "irrigation_plan": irrig,
            "equipment_list": eqr,
            "knowledge_refs": krefs,
        })
        log_json("5-编排结果", {
            "content_length": len(orch_result["content"]),
            "sources_count": len(krefs),
        })

        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "quality", "message": "正在进行方案质量审核评估..."})

        quality = QualityAgent(db, knowledge_forest)
        qr = await quality.check(orch_result["content"], profile)
        log_json("6-质量审核", {
            "compliance": qr["compliance"]["score"],
            "rationality": qr["rationality"]["score"],
            "sufficiency": qr["sufficiency"]["score"],
            "overall": qr["overall_score"],
            "grade": qr["grade"],
        })

        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "llm", "message": "正在调用AI模型优化方案内容..."})

        if deepseek._loaded:
            design_knowledge = irrig.get("knowledge_learned", [])
            equip_knowledge = eqr.get("knowledge_sources", [])
            prompt = build_irrigation_prompt(
                profile, krefs, "",
                design_knowledge, equip_knowledge,
                template_refs=template_refs,
                tech_refs=tech_refs,
                standard_refs=standard_refs,
            )
            llm_gen = await deepseek.generate_async(prompt, max_tokens=4096)
            if llm_gen:
                final_content = llm_gen
            else:
                final_content = orch_result["content"]
        else:
            final_content = orch_result["content"]

        crop_inferred = profile.get("crop_inferred", False)
        crop_note = "（系统根据地区自动推断）" if crop_inferred else ""

        summary = f"""{profile.get('location','')}{profile.get('crop_type','')}灌溉方案已生成{crop_note}。

| 项目 | 数值 |
|------|------|
| 作物 | {profile.get('crop_type','')}{crop_note} |
| 面积 | {profile.get('area','')}亩 |
| 预算 | {budget:,.0f}元 |
| 灌溉方式 | {irrig.get('project_info',{}).get('irrigation_method','')} |
| 总需水量 | {irrig.get('project_info',{}).get('total_water_requirement_m3','')}m³ |
| 设备投资 | {eqr.get('total_cost',''):,.0f}元 |
| 质量评分 | {qr['overall_score']}分({qr['grade']}) |
| 溯源参考 | {len(krefs)}条知识依据(模板{sum(1 for r in krefs if r.get('channel')=='template')}/技术{sum(1 for r in krefs if r.get('channel')=='tech')}/标准{sum(1 for r in krefs if r.get('channel')=='standard')}) |
| LLM模式 | {'本地模型' if deepseek.mode == 'local' else '云端API' if deepseek.mode == 'api' else '规则引擎'} |"""

        ctx_history.append({"role": "assistant", "content": final_content})
        _save_msg_to_db(db, sid, "assistant", final_content)
        if sid in conversation_context:
            conversation_context[sid] = ctx_history

        yield _sse({
            "type": "complete",
            "session_id": sid,
            "reply": final_content,
            "data": {
                "summary": summary,
                "plan_content": final_content,
                "quality": {
                    "compliance": qr["compliance"],
                    "rationality": qr["rationality"],
                    "sufficiency": qr["sufficiency"],
                    "overall_score": qr["overall_score"],
                    "grade": qr["grade"],
                    "suggestions": qr["suggestions"],
                },
                "profile": {k: profile[k] for k in ["location", "crop_type", "area", "budget"]},
                "krefs": [{"source": r["source"], "similarity": r["similarity"]} for r in krefs[:5]],
                "knowledge_refs": [
                    {
                        "source": r.get("source", "") or r.get("chunk_id", "") or "未知文献",
                        "content": r.get("content", "")[:300],
                        "channel": r.get("channel", ""),
                        "similarity": round(float(r.get("similarity", 0)), 4),
                        "tag": r.get("tag", ""),
                    }
                    for r in krefs
                ],
                "llm_enabled": deepseek._loaded,
                "llm_mode": deepseek.mode,
            },
        })
        log_json("COMPLETE_SENT", {
            "knowledge_refs_count": len(krefs),
            "knowledge_refs_sample": [{"source": r.get("source", ""), "channel": r.get("channel", "")} for r in krefs[:3]],
        })

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@app.post("/chat/modify")
async def chat_modify_stream(msg: ChatMsg, db: Session = Depends(get_db)):
    global knowledge_forest, conversation_context
    sid = msg.session_id or str(uuid.uuid4())
    ui = msg.message.strip()
    original_plan = msg.original_plan if hasattr(msg, "original_plan") else None

    if sid not in conversation_context:
        conversation_context[sid] = _load_history_from_db(db, sid)
    ctx_history = conversation_context[sid]

    if not original_plan:
        for msg_obj in reversed(ctx_history):
            if msg_obj.get("role") == "assistant" and len(msg_obj.get("content", "")) > 500 and any(ch in msg_obj.get("content", "") for ch in ["一、", "1. 项目概况", "二、", "三、"]):
                original_plan = msg_obj["content"]
                break

    async def event_stream() -> AsyncGenerator[str, None]:
        yield _sse({"type": "progress", "step": "start", "message": "开始局部修改方案..."})

        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "analyze", "message": "步骤1: 解析修改意图，识别影响模块..."})

        from backend.agents.modify import PlanModifyAgent
        from backend.agents.modify import MODULE_ORDER as PLAN_MODULES
        modifier = PlanModifyAgent(db)

        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "modify", "message": "步骤2-4: 依赖分析 + 局部检索 + 分块生成..."})

        modified_result = await modifier.modify(original_plan, ui)
        final_content = modified_result["content"]
        quality_report = modified_result["quality"]
        affected_modules = modified_result.get("affected_sections", [])
        warnings = modified_result.get("warnings", [])
        diff_data = modified_result.get("diff", {})
        knowledge_refs = modified_result.get("knowledge_refs", [])
        retry_info = modified_result.get("retry_info", {})

        # 使用 QualityAgent 对全文进行完整质量审核（覆盖局部审核）
        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "quality", "message": "正在进行全文质量审核评估..."})
        profile = {}
        for msg in ctx_history:
            if msg["role"] == "user":
                extracted = RequirementAgent._raw_extract_entities(msg["content"])
                profile.update({k: v for k, v in extracted.items() if v})
        quality = QualityAgent(db, knowledge_forest)
        qr = await quality.check(final_content, profile)

        logger.info(f"[质检] 局部审核分数: {quality_report.get('overall_score', 'N/A')}")
        logger.info(f"[质检] 全文审核分数: {qr.get('overall_score', 'N/A')}")

        # 使用全文审核结果（更准确）
        quality_report = {
            "overall_score": qr.get("overall_score", 0),
            "grade": qr.get("grade", "合格"),
            "compliance": qr.get("compliance", {}),
            "rationality": qr.get("rationality", {}),
            "sufficiency": qr.get("sufficiency", {}),
            "suggestions": qr.get("suggestions", []),
        }

        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "consistency", "message": f"步骤5-6: 一致性检查完成，{len(warnings)}条警告"})

        await asyncio.sleep(0.1)
        yield _sse({"type": "progress", "step": "complete", "message": f"局部修改完成（影响{len(affected_modules)}个模块）"})

        ctx_history.append({"role": "assistant", "content": final_content})
        _save_msg_to_db(db, sid, "assistant", final_content)
        if sid in conversation_context:
            conversation_context[sid] = ctx_history

        yield _sse({
            "type": "complete",
            "session_id": sid,
            "reply": final_content,
            "data": {
                "quality": quality_report,
                "affected_modules": [m for m in affected_modules if m in PLAN_MODULES],
                "warnings": warnings,
                "diff": diff_data,
                "knowledge_refs": knowledge_refs,
                "retry_info": retry_info,
            },
        })

    return StreamingResponse(event_stream(), media_type="text/event-stream")


def _sse(data: dict) -> str:
    return f"data: {json.dumps(data, ensure_ascii=False)}\n\n"


def _extract_location(text: str) -> str:
    for name in sorted(CHINA_COORDS.keys(), key=len, reverse=True):
        if name in text:
            return name
    return ""

async def _do_weather(sid: str, ui: str):
    loc = _extract_location(ui)
    # Priority 1: session_entities (user-provided info)
    if not loc and sid in session_entities and session_entities[sid].get("location"):
        loc = session_entities[sid]["location"]
        logger.info(f"[气象] 从session_entities提取: {loc}")
    # Priority 2: search conversation context (avoid LLM-generated content mentioning other cities)
    if not loc and sid in conversation_context:
        ctx = conversation_context[sid]
        for idx, m in enumerate(reversed(ctx)):
            if m.get("role") == "user":
                found = _extract_location(m.get("content", ""))
                if found:
                    loc = found
                    logger.info(f"[气象] 从用户消息提取: {found}")
                    break
    loc = loc or "北京"
    da = DataAcquisition()
    weather = await da.get_weather_data(loc)
    log_json("气象数据", {"location": loc, "data": weather})
    reply = f"""## {loc} 实时气象数据

| 指标 | 数值 |
|------|------|
| 气温 | {weather['temperature']}°C |
| 湿度 | {weather['humidity']}% |
| 风速 | {weather['wind_speed']} m/s |
| 参考ET0 | {weather['et0']} mm/天 |
| 数据源 | {weather['source']} |"""

    if sid in conversation_context:
        conversation_context[sid].append({"role": "assistant", "content": reply})
    return {"success": True, "session_id": sid, "reply": reply, "type": "weather", "data": weather}


async def _do_soil(sid: str, ui: str):
    loc = _extract_location(ui)
    if not loc and sid in session_entities and session_entities[sid].get("location"):
        loc = session_entities[sid]["location"]
        logger.info(f"[土壤] 从session_entities提取: {loc}")
    if not loc and sid in conversation_context:
        for m in reversed(conversation_context[sid]):
            found = _extract_location(m.get("content", ""))
            if found:
                loc = found
                logger.info(f"[土壤] 从上下文提取: {found}")
                break
    loc = loc or "北京"
    da = DataAcquisition()
    soil = await da.get_soil_data(loc)
    log_json("土壤数据", {"location": loc, "data": soil})
    reply = f"""## {loc} 土壤数据

| 指标 | 数值 |
|------|------|
| 土壤类型 | {soil['soil_type']} |
| 土壤温度 | {soil['soil_temperature']}°C |
| 田间持水量 | {soil['field_capacity']:.1%} |
| 萎蔫系数 | {soil['wilting_point']:.1%} |
| 数据源 | {soil['source']} |"""

    if sid in conversation_context:
        conversation_context[sid].append({"role": "assistant", "content": reply})
    return {"success": True, "session_id": sid, "reply": reply, "type": "soil", "data": soil}


async def _do_knowledge_browse(db: Session, sid: str):
    global knowledge_forest
    if knowledge_forest is None:
        knowledge_forest = KnowledgeForest(db)
    if not knowledge_forest.is_built:
        if not knowledge_forest.try_load():
            await knowledge_forest.build_forest()

    overview = knowledge_forest.get_doc_overview()
    log_json("知识森林浏览", overview)

    from backend.knowledge_forest.tags import get_tag_label

    lines = ["## 知识森林全貌\n"]
    lines.append(f"> **总节点**: {overview['total_nodes']} 个")
    lines.append(f"> **文档来源**: {overview['total_sources']} 个\n")

    lines.append(f"### 层次结构")
    lines.append(f"| 层级 | 节点数 | 说明 |")
    lines.append(f"|------|--------|------|")
    lines.append(f"| L0（文本块） | {overview['L0_total']} | 原始文档切分后的文本片段 |")
    lines.append(f"| L1（单文档摘要） | {overview['L1_total']} | 每篇文档内聚类生成的摘要 |")
    lines.append(f"| L2（跨文档摘要） | {overview['L2_total']} | 跨文档聚类生成的高层知识 |\n")

    lines.append(f"### 文档清单（{overview['total_sources']} 篇）")
    lines.append(f"| # | 类型 | 文档名称 | 节点数 |")
    lines.append(f"|---|------|----------|--------|")
    for i, doc in enumerate(overview["documents"], 1):
        tag_label = doc.get("tag_label") or get_tag_label(doc["tag"])
        lines.append(f"| {i} | {tag_label} | {doc['source']} | {doc['total_nodes']} |")

    lines.append(f"\n> 输入具体关键词（如\"灌溉制度\"、\"FAO\"）即可检索知识森林中的相关内容。")

    reply = "\n".join(lines)
    if sid in conversation_context:
        conversation_context[sid].append({"role": "assistant", "content": reply})
    return {"success": True, "session_id": sid, "reply": reply, "type": "knowledge_browse", "data": overview}


async def _do_knowledge(db: Session, sid: str, ui: str):
    global knowledge_forest
    if knowledge_forest is None:
        knowledge_forest = KnowledgeForest(db)
    if not knowledge_forest.is_built:
        if not knowledge_forest.try_load():
            await knowledge_forest.build_forest()

    results = await knowledge_forest.query(ui, k=5)
    log_json("知识检索", {"query": ui, "results": results})

    lines = ["## 知识森林检索结果\n"]

    if results:
        for i, r in enumerate(results, 1):
            lines.append(f"**{i}.** `{r['source']}` (相似度: {r['similarity']:.2%})")
            lines.append(f"> {r['content'][:200]}...\n")
    else:
        lines.append("> 未找到匹配的知识条目，请尝试更具体的关键词。")

    reply = "\n".join(lines)
    if sid in conversation_context:
        conversation_context[sid].append({"role": "assistant", "content": reply})
    return {"success": True, "session_id": sid, "reply": reply, "type": "knowledge", "data": results}


@app.get("/conversation/{session_id}")
async def get_conversation(session_id: str, db: Session = Depends(get_db)):
    global conversation_context
    db_history = _load_history_from_db(db, session_id)
    mem_history = conversation_context.get(session_id, [])
    merged = db_history + [m for m in mem_history if m not in db_history]
    return {"session_id": session_id, "history": merged, "count": len(merged)}

@app.post("/knowledge/build")
async def build_knowledge(db: Session = Depends(get_db)):
    global knowledge_forest
    knowledge_forest = KnowledgeForest(db)
    result = await knowledge_forest.build_forest()
    log_json("知识森林构建完成", result)
    return {"success": True, "result": result}


@app.get("/knowledge/status")
async def knowledge_status(db: Session = Depends(get_db)):
    global knowledge_forest
    if knowledge_forest is None:
        knowledge_forest = KnowledgeForest(db)
    return knowledge_forest.get_status()


@app.post("/query-knowledge")
async def query_knowledge(query: str, tag: str = None, db: Session = Depends(get_db)):
    global knowledge_forest
    if knowledge_forest is None:
        knowledge_forest = KnowledgeForest(db)
    if not knowledge_forest.is_built:
        if not knowledge_forest.try_load():
            await knowledge_forest.build_forest()

    if tag:
        results = await knowledge_forest.tagged_query(query, tag, k=10)
        log_json("标签知识查询", {"query": query, "tag": tag, "results_count": len(results)})
    else:
        results = await knowledge_forest.query(query, k=10)
        log_json("知识查询", {"query": query, "results": results})
    return results


@app.post("/acquire-weather")
async def acquire_weather(location: str = "北京"):
    da = DataAcquisition()
    return await da.get_weather_data(location)


@app.post("/acquire-soil")
async def acquire_soil(location: str = "北京"):
    da = DataAcquisition()
    return await da.get_soil_data(location)


@app.get("/llm/status")
async def llm_status():
    return {
        "loaded": deepseek._loaded,
        "mode": deepseek.mode,
        "api_model": deepseek.api_model,
        "api_base": deepseek.api_base,
    }

@app.get("/export/{session_id}")
async def export_plan(session_id: str, fmt: str = "md"):
    global conversation_context
    ctx = conversation_context.get(session_id, [])

    plan_content = ""
    for m in ctx:
        if m["role"] == "assistant" and len(m["content"]) > 500 and any(ch in m["content"] for ch in ["一、", "1. 项目概况", "二、", "三、"]):
            plan_content = m["content"]
            break

    if not plan_content:
        raise HTTPException(404, "未找到方案内容")

    if fmt == "md":
        return Response(
            content=plan_content,
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename=irrigation_{session_id[:8]}.md"},
        )
    elif fmt == "docx":
        try:
            from docx import Document
            doc = Document()
            for line in plan_content.split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("|"):
                    continue
                elif line.strip():
                    doc.add_paragraph(line.strip())

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return Response(
                content=buf.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=irrigation_{session_id[:8]}.docx"},
            )
        except ImportError:
            json_content = json.dumps({"content": plan_content}, ensure_ascii=False)
            return Response(
                content=json_content,
                media_type="application/json",
                headers={"Content-Disposition": f"attachment; filename=irrigation_{session_id[:8]}.json"},
            )
    elif fmt == "pdf":
        json_content = json.dumps({"content": plan_content}, ensure_ascii=False)
        return Response(
            content=json_content,
            media_type="application/json",
            headers={"Content-Disposition": f"attachment; filename=irrigation_{session_id[:8]}.json"},
        )
    elif fmt == "txt":
        return Response(
            content=plan_content,
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename=irrigation_{session_id[:8]}.txt"},
        )
    else:
        raise HTTPException(400, f"不支持的格式: {fmt}，支持: md, docx, txt")


@app.get("/greeting")
async def greeting():
    return {
        "reply": "你好！我是智慧灌溉方案生成助手 。\n\n你可以：\n1. 描述项目，如\"河南中牟县2000亩小麦，预算200万\"，我来生成方案\n2. 不指定作物时，系统会根据地区自动推断最适合的作物\n3. 查询气象/土壤数据\n4. 查询知识库\n5. 对已有方案提出修改"
    }


frontend_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "frontend")
if os.path.exists(frontend_path):
    @app.get("/")
    async def serve_frontend():
        return FileResponse(os.path.join(frontend_path, "index.html"))

    @app.get("/lib/{filename}")
    async def serve_lib(filename: str):
        file_path = os.path.join(frontend_path, "lib", filename)
        if os.path.exists(file_path):
            return FileResponse(file_path)
        return {"error": f"File not found: {filename}"}, 404


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)